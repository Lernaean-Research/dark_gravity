from __future__ import annotations

import re
from pathlib import Path

from docx import Document


def main() -> int:
    path = Path("Kitcey_2026_Response_Sector_v5_0_1_1_WITH_APPENDIX_C.docx")
    doc = Document(str(path))

    texts = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]

    sentinel = "Appendix C. Cluster morphology operator and HFF benchmark"
    print("sentinel_found:", any(sentinel in t for t in texts))

    fig_caps = [t for t in texts if t.startswith("Figure C")]
    print("figure_captions:", fig_caps)

    pat = re.compile(r"^\[(\d+)\]")
    nums = [int(m.group(1)) for t in texts for m in [pat.match(t)] if m]
    print("max_ref:", max(nums) if nums else None)

    refs = [t for t in texts if pat.match(t)]
    print("last_refs:")
    for t in refs[-8:]:
        print(" ", t[:160])

    # Show context around a References heading to infer formatting.
    ref_heading_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text and p.text.strip().lower() in {"references", "bibliography"}:
            ref_heading_idx = i
            break
    print("ref_heading_found:", ref_heading_idx is not None)
    if ref_heading_idx is not None:
        print("refs_context:")
        for p in doc.paragraphs[ref_heading_idx : ref_heading_idx + 25]:
            t = (p.text or "").strip()
            if not t:
                continue
            print(" ", t[:180])

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
