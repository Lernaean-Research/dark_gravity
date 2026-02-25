"""Insert Appendix C (HFF cluster morphology operator) into the main DOCX.

Targets:
- Input:  Spacetime_Mechanics/Kitcey_2026_Response_Sector_v5_0_1_1.docx
- Output: Spacetime_Mechanics/Kitcey_2026_Response_Sector_v5_0_1_1_WITH_APPENDIX_C.docx

Design goals:
- Do not duplicate if run multiple times (sentinel heading check).
- Insert Appendix C before the References section when possible.
- Insert new reference entries into the existing References list (author–date).

This script intentionally keeps formatting simple and Word-friendly.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
IN_DOCX = ROOT / "Kitcey_2026_Response_Sector_v5_0_1_1.docx"
OUT_DOCX = ROOT / "Kitcey_2026_Response_Sector_v5_0_1_1_WITH_APPENDIX_C.docx"

SUMMARY_FIG = ROOT / "toy_models" / "out_predictions" / "figures" / "hff_systematics_summary_roi100.png"
A2744_EXAMPLE = (
    ROOT
    / "toy_models"
    / "out_predictions"
    / "figures"
    / "systematics_sixpanel"
    / "abell2744"
    / "abell2744_cats_v4.1_roi100_sixpanel.png"
)
MACS_EXAMPLE = (
    ROOT
    / "toy_models"
    / "out_predictions"
    / "figures"
    / "systematics_sixpanel"
    / "macs0416"
    / "macs0416_cats_v4.1_roi100_sixpanel.png"
)

SENTINEL = "Appendix C. Cluster morphology operator and HFF benchmark"


def _iter_paragraph_texts(doc: Document):
    for p in doc.paragraphs:
        yield p.text.strip()


def _has_appendix_already(doc: Document) -> bool:
    return any(SENTINEL.lower() in t.lower() for t in _iter_paragraph_texts(doc))


def _find_paragraph_index(doc: Document, predicate) -> int | None:
    for i, p in enumerate(doc.paragraphs):
        if predicate(p):
            return i
    return None


def _extract_max_reference_number(doc: Document) -> int:
    # Look for paragraphs starting with [N]
    pattern = re.compile(r"^\s*\[(\d+)\]\s+")
    max_n = 0
    for t in _iter_paragraph_texts(doc):
        m = pattern.match(t)
        if m:
            max_n = max(max_n, int(m.group(1)))
    return max_n


def _find_references_heading(doc: Document) -> Paragraph | None:
    for p in doc.paragraphs:
        if p.text and p.text.strip().lower() in {"references", "bibliography"}:
            return p
    return None


def _reference_key(text: str) -> str:
    t = text.strip()
    if not t:
        return ""
    head = re.split(r"[,.(]", t, maxsplit=1)[0].strip()
    return head.casefold()


def _has_reference_like(doc: Document, needle: str) -> bool:
    n = needle.casefold()
    return any(n in (p.text or "").casefold() for p in doc.paragraphs)


def _insert_reference_apa_sorted(doc: Document, ref_text: str) -> None:
    """Insert `ref_text` into the References section in alphabetical order.

    If the document has no References heading, appends at end.
    If a very similar reference is already present, does nothing.
    """

    doi_match = re.search(r"10\.\d{4,9}/[^\s)]+", ref_text)
    if doi_match and _has_reference_like(doc, doi_match.group(0)):
        return
    if _has_reference_like(doc, ref_text[:40]):
        return

    refs_heading = _find_references_heading(doc)
    if refs_heading is None:
        doc.add_paragraph(ref_text)
        return

    # Find heading index.
    start_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p is refs_heading:
            start_idx = i
            break
    if start_idx is None:
        doc.add_paragraph(ref_text)
        return

    new_key = _reference_key(ref_text)
    ref_paras = [p for p in doc.paragraphs[start_idx + 1 :] if (p.text or "").strip()]
    if not ref_paras:
        _insert_paragraph_after(refs_heading, ref_text)
        return

    for p in ref_paras:
        k = _reference_key(p.text or "")
        if k and k > new_key:
            p.insert_paragraph_before(ref_text)
            return

    _insert_paragraph_after(ref_paras[-1], ref_text)


def _insert_paragraph_after(paragraph: Paragraph, text: str, style: str | None = None) -> Paragraph:
    """Insert a new paragraph directly after `paragraph`."""

    new_p_elm = OxmlElement("w:p")
    paragraph._p.addnext(new_p_elm)
    new_p = Paragraph(new_p_elm, paragraph._parent)

    if text:
        new_p.add_run(text)

    if style:
        try:
            new_p.style = style
        except Exception:
            pass

    return new_p


def _insert_picture_after(paragraph: Paragraph, image_path: Path, width_inches: float) -> Paragraph:
    p = _insert_paragraph_after(paragraph, "")
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    return p


def _safe_heading_style(doc: Document, level: int) -> str | None:
    # Word's default heading styles are 'Heading 1', 'Heading 2', ...
    candidate = f"Heading {level}"
    try:
        _ = doc.styles[candidate]
        return candidate
    except Exception:
        return None


def main() -> int:
    if not IN_DOCX.exists():
        raise FileNotFoundError(IN_DOCX)

    for p in [SUMMARY_FIG, A2744_EXAMPLE, MACS_EXAMPLE]:
        if not p.exists():
            raise FileNotFoundError(f"Missing figure: {p}")

    doc = Document(str(IN_DOCX))

    if _has_appendix_already(doc):
        print("Appendix C already present; writing a copy anyway.")
        doc.save(str(OUT_DOCX))
        print(f"Wrote: {OUT_DOCX}")
        return 0

    # Try to insert before the References heading.
    ref_idx = _find_paragraph_index(
        doc,
        lambda p: p.text.strip().lower() in {"references", "bibliography"},
    )

    insert_after_idx = (ref_idx - 1) if (ref_idx is not None and ref_idx > 0) else (len(doc.paragraphs) - 1)
    anchor = doc.paragraphs[insert_after_idx]

    h1 = _safe_heading_style(doc, 1)
    h2 = _safe_heading_style(doc, 2)

    last_p = anchor

    # Appendix heading
    last_p = _insert_paragraph_after(last_p, SENTINEL, style=h1)

    # Body text (kept close to the MD, but Word-friendly)
    body_blocks = [
        (
            "This appendix documents the preregistered cluster morphology operator used to compare "
            "a gravitational-lensing convergence field (κ; treated here as a response proxy) against "
            "an observational hot-gas tracer constructed from Chandra imaging. The benchmark dataset "
            "is two Hubble Frontier Fields (HFF) clusters (Abell 2744 and MACS J0416.1−2403) using "
            "public HFF lens-model products (Mikulski Archive for Space Telescopes [MAST], 2019)."
        ),
        "C.1 Data provenance and scope",
        "• κ maps (multi-team lens models): HFF community lens reconstructions accessed from the MAST Frontier Fields HLSP archive (MAST, 2019).",
        "• X-ray proxy (stacked Chandra img2 maps): HEASARC-served *_full_img2.fits.gz image products were stacked into a common WCS grid and used as a morphology/centroid proxy for the hot intracluster medium (ICM) (NASA/GSFC HEASARC, 2025; Weisskopf et al., 2002).",
        (
            "For reproducibility, the X-ray proxy construction follows a lightweight procedure implemented in "
            "toy_models/make_chandra_xray_map.py: each img2 image is divided by a scalar exposure keyword "
            "(EXPOSURE, with LIVETIME/ONTIME fallbacks), reprojected onto a common WCS grid, and combined "
            "as an exposure-weighted mean rate map."
        ),
        (
            "Important limitation: this work does not perform event-level Chandra reduction (exposure map correction, "
            "background modeling, point-source masking, etc.) in CIAO [23]. The X-ray product is therefore intended "
            "as a geometric locator of bright ICM structure for centroid comparisons, not a calibrated surface-brightness "
            "measurement."
        ),
        "C.2 Preregistered morphology operator",
        "For each cluster, each κ team model, and each ROI radius, we apply the same fixed operator to κ and to the X-ray proxy map:",
        "1) Smoothing: apply Gaussian smoothing with σ = 8 arcsec independently to κ and X-ray maps.",
        "2) ROI restriction: within a circular ROI (fixed ICRS center per cluster), compute pixel-intensity thresholds at the 99th, 97th, and 95th percentiles of the smoothed pixels.",
        "3) Primary-blob selection: for each thresholded map, select the largest connected component above threshold.",
        "4) Centroid definition: compute the unweighted mask centroid of that connected component.",
        "5) Measured quantity: record the κ–X-ray centroid separation in arcseconds.",
        (
            "This definition is intentionally simple and falsifiable: it is fixed across clusters, teams, and thresholds, "
            "with only the ROI radius swept to quantify footprint sensitivity."
        ),
        "C.3 Robustness grid: multi-team systematics and ROI sensitivity",
        (
            "We evaluate ROI radius ∈ {80″, 100″, 120″} for each cluster across available HFF teams, and report both "
            "(i) cross-team spread at fixed ROI, and (ii) ROI sensitivity within each team model. The complete tables "
            "are in toy_models/HFF_ALL_TEAMS_SYSTEMATICS_ANALYSIS.md and the per-team measurements are in "
            "toy_models/out_predictions/systematics/{abell2744,macs0416}/systematics_summary.csv."
        ),
        (
            "At ROI = 100″, Abell 2744 shows a relatively stable median κ–X-ray separation (≈ 66–67″ across thresholds) "
            "with small cross-team IQRs (≈ 1–1.5″) but with a nontrivial full range at high thresholds. "
            "MACS J0416.1−2403 shows much larger cross-team dispersion at higher thresholds (IQRs growing to tens of arcseconds), "
            "reflecting stronger lens-model dependence of the inferred morphology under this operator."
        ),
        "Figure C1: HFF κ–X-ray systematics summary at ROI = 100″.",
        "C.4 Six-panel diagnostic figures (ROI = 100″)",
        (
            "For interpretability, we generated per-team six-panel diagnostics at ROI = 100″, including κ, X-ray proxy, "
            "threshold masks, and centroid overlays. Representative examples are shown below; the complete figure sets are "
            "available under toy_models/out_predictions/figures/systematics_sixpanel/."
        ),
        "Figure C2: Abell 2744 (CATS v4.1), ROI = 100″.",
        "Figure C3: MACS J0416.1−2403 (CATS v4.1), ROI = 100″.",
    ]

    # Convert any remaining bracketed CIAO citation to author-date.
    body_blocks = [
        b.replace("in CIAO [23]", "in CIAO (Fruscione et al., 2006)") if isinstance(b, str) else b
        for b in body_blocks
    ]

    # Insert blocks; treat section labels as Heading 2 when possible.
    for block in body_blocks:
        style = None
        if block in {"C.1 Data provenance and scope", "C.2 Preregistered morphology operator", "C.3 Robustness grid: multi-team systematics and ROI sensitivity", "C.4 Six-panel diagnostic figures (ROI = 100\")"}:
            style = h2
        last_p = _insert_paragraph_after(last_p, block, style=style)

        # Insert pictures after the relevant figure captions.
        if block.startswith("Figure C1"):
            last_p = _insert_picture_after(last_p, SUMMARY_FIG, width_inches=6.5)
        if block.startswith("Figure C2"):
            last_p = _insert_picture_after(last_p, A2744_EXAMPLE, width_inches=6.5)
        if block.startswith("Figure C3"):
            last_p = _insert_picture_after(last_p, MACS_EXAMPLE, width_inches=6.5)

    # Insert APA-style references into the document's References list.
    new_refs_apa = [
        (
            "Fruscione, A., et al. (2006). CIAO: Chandra’s Data Analysis System. "
            "Proceedings of SPIE, 6270, 62701V. https://cxc.cfa.harvard.edu/ciao/"
        ),
        (
            "Mikulski Archive for Space Telescopes (MAST). (2019). Hubble Space Telescope Frontier Fields (HLSP archive). "
            "Retrieved February 24, 2026, from https://archive.stsci.edu/hlsp/frontier/"
        ),
        (
            "NASA/GSFC HEASARC. (2025). Archive access and documentation. "
            "Retrieved February 24, 2026, from https://heasarc.gsfc.nasa.gov/docs/archive.html"
        ),
        (
            "Weisskopf, M. C., et al. (2002). An Overview of the Performance and Scientific Results from the Chandra X-Ray Observatory. "
            "Publications of the Astronomical Society of the Pacific, 114, 1. https://doi.org/10.1086/338108"
        ),
    ]
    for ref in new_refs_apa:
        _insert_reference_apa_sorted(doc, ref)

    doc.save(str(OUT_DOCX))
    print(f"Wrote: {OUT_DOCX}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
