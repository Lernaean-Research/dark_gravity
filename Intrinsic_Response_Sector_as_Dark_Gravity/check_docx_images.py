from docx import Document
import os

doc = Document(r'D:\#Documents\#Publication\Spacetime_Mechanics\Dark Gravity dominated stellar clusters\Kitcey_2026_Dark_Gravity_Dominated_Galaxies_and_Globular_Clusters.v.2.0.docx')

# Get image relationships with sizes
rels = doc.part.rels
img_rels_dict = {r_id: (rels[r_id], len(rels[r_id].target_part.blob)) 
                 for r_id in rels if 'image' in rels[r_id].target_ref}

print("=" * 80)
print("IMAGES IN ORDER OF APPEARANCE:")
print("=" * 80)

figure_num = 0
for i, para in enumerate(doc.paragraphs):
    # Check if this paragraph has an image
    for run in para.runs:
        # Check for drawing elements (images)
        drawings = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
        for drawing in drawings:
            # Get the blip element which contains the image reference
            blips = drawing.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
            for blip in blips:
                r_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if r_id in img_rels_dict:
                    rel, size = img_rels_dict[r_id]
                    name = os.path.basename(rel.target_ref)
                    figure_num += 1
                    print(f"\nFigure {figure_num}: {name} ({size:,} bytes)")
                    print(f"   Para index: {i}")
                    # Look for caption in next paragraph
                    if i + 1 < len(doc.paragraphs):
                        caption = doc.paragraphs[i+1].text.strip()
                        if caption.startswith('Figure'):
                            print(f"   Caption: {caption[:150]}")

print("\n" + "=" * 80)
print("MAPPING TO SOURCE FILES:")
print("=" * 80)

source_files = {
    'Picture1.jpg': 62105,
    'image (1).jpg': 338019,
    'image (2).jpg': 160368,
    'image (3).jpg': 270316,
    'image (4).jpg': 393399,
    'image.jpg': 873659
}

print("\nDOCX -> Source file mapping:")
for name, size in [('image1.jpg', 338019), ('image2.jpg', 160368), ('image3.jpg', 873659), ('image4.jpg', 393399)]:
    matching = [f for f, s in source_files.items() if s == size]
    if matching:
        print(f"  {name} ({size:,} bytes) -> {matching[0]}")
    else:
        print(f"  {name} ({size:,} bytes) -> NO MATCH in source folder!")
