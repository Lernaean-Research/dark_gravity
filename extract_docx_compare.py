"""
Extract content from Word document and prepare for comparison with LaTeX manuscript.
"""
import docx
import re
import json
from pathlib import Path
from collections import defaultdict

def extract_docx_content(docx_path):
    """
    Extract structured content from Word document.
    Returns dict with sections, paragraphs, formatting info, etc.
    """
    doc = docx.Document(docx_path)
    
    extracted = {
        'full_text': [],
        'sections': [],
        'paragraphs': [],
        'tables': [],
        'metadata': {
            'total_paragraphs': 0,
            'total_words': 0,
            'total_chars': 0,
        }
    }
    
    current_section = None
    section_stack = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
            
        # Parse paragraph formatting
        style_name = para.style.name if para.style else 'Normal'
        
        # Check for bold/italic/formatting
        runs_info = []
        for run in para.runs:
            if run.text.strip():
                runs_info.append({
                    'text': run.text,
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                })
        
        para_info = {
            'index': i,
            'text': text,
            'style': style_name,
            'runs': runs_info,
            'word_count': len(text.split()),
        }
        
        # Detect section headers (heuristic based on style or formatting)
        is_heading = 'Heading' in style_name or style_name.startswith('Title')
        
        if is_heading:
            level = 1
            if 'Heading 1' in style_name or 'Title' in style_name:
                level = 1
            elif 'Heading 2' in style_name:
                level = 2
            elif 'Heading 3' in style_name:
                level = 3
            else:
                # Try to extract from style name
                match = re.search(r'Heading (\d+)', style_name)
                if match:
                    level = int(match.group(1))
            
            section_info = {
                'title': text,
                'level': level,
                'para_index': i,
                'content': []
            }
            
            # Manage section stack
            while section_stack and section_stack[-1]['level'] >= level:
                section_stack.pop()
            
            if section_stack:
                section_stack[-1]['subsections'] = section_stack[-1].get('subsections', [])
                section_stack[-1]['subsections'].append(section_info)
            else:
                extracted['sections'].append(section_info)
            
            section_stack.append(section_info)
            current_section = section_info
        else:
            # Regular content paragraph
            if current_section is not None:
                current_section['content'].append(para_info)
        
        extracted['paragraphs'].append(para_info)
        extracted['full_text'].append(text)
        extracted['metadata']['total_words'] += para_info['word_count']
        extracted['metadata']['total_chars'] += len(text)
    
    extracted['metadata']['total_paragraphs'] = len(extracted['paragraphs'])
    
    # Extract tables
    for table_idx, table in enumerate(doc.tables):
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        extracted['tables'].append({
            'index': table_idx,
            'data': table_data,
            'rows': len(table_data),
            'cols': len(table_data[0]) if table_data else 0,
        })
    
    return extracted

def extract_latex_structure(tex_path):
    """
    Parse LaTeX file to extract structure for comparison.
    """
    with open(tex_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    extracted = {
        'full_text': content,
        'sections': [],
        'paragraphs': [],
        'citations': [],
        'metadata': {
            'total_lines': content.count('\n'),
            'total_words': 0,
            'total_chars': len(content),
        }
    }
    
    # Find all sections/subsections
    section_pattern = r'\\(section|subsection|subsubsection)\{([^}]+)\}'
    for match in re.finditer(section_pattern, content):
        section_type = match.group(1)
        section_title = match.group(2)
        level = {'section': 1, 'subsection': 2, 'subsubsection': 3}.get(section_type, 1)
        
        extracted['sections'].append({
            'type': section_type,
            'title': section_title,
            'level': level,
            'position': match.start(),
        })
    
    # Find citations
    citation_pattern = r'\\cite[tp]?\{([^}]+)\}'
    for match in re.finditer(citation_pattern, content):
        citations = match.group(1).split(',')
        for cite in citations:
            extracted['citations'].append(cite.strip())
    
    # Count words (approximate - remove LaTeX commands)
    text_only = re.sub(r'\\[a-zA-Z]+(\[[^\]]*\])?(\{[^}]*\})*', '', content)
    text_only = re.sub(r'[{}%$\\]', '', text_only)
    words = text_only.split()
    extracted['metadata']['total_words'] = len(words)
    
    return extracted

def compare_documents(docx_data, latex_data):
    """
    Compare extracted data and identify differences.
    """
    comparison = {
        'statistics': {},
        'new_sections': [],
        'removed_sections': [],
        'new_citations': [],
        'removed_citations': [],
        'text_differences': []
    }
    
    # Statistics comparison
    comparison['statistics'] = {
        'docx': {
            'words': docx_data['metadata']['total_words'],
            'sections': len(docx_data['sections']),
            'paragraphs': docx_data['metadata']['total_paragraphs'],
            'tables': len(docx_data['tables']),
        },
        'latex': {
            'words': latex_data['metadata']['total_words'],
            'sections': len(latex_data['sections']),
            'lines': latex_data['metadata']['total_lines'],
            'citations': len(set(latex_data['citations'])),
        }
    }
    
    # Section comparison
    docx_section_titles = set(s['title'] for s in docx_data['sections'])
    latex_section_titles = set(s['title'] for s in latex_data['sections'])
    
    comparison['new_sections'] = list(docx_section_titles - latex_section_titles)
    comparison['removed_sections'] = list(latex_section_titles - docx_section_titles)
    
    # Citation comparison (if we can extract from docx)
    docx_text = ' '.join(docx_data['full_text'])
    docx_citations = set(re.findall(r'\[([^\]]+)\]', docx_text))
    latex_citations = set(latex_data['citations'])
    
    comparison['new_citations'] = list(docx_citations - latex_citations)
    
    return comparison

def main():
    # Paths
    docx_path = r"D:\#Documents\#Publication\Spacetime_Mechanics__git\Kitcey_2026_Intrinsic_Response_Sector_DM_Candidacy.v5.1.docx"
    latex_path = r"d:\#Documents\#Publication\Spacetime_Mechanics__git\Intrinsic_Response_Sector_as_Dark_Gravity\manuscript_overleaf.tex"
    output_path = r"d:\#Documents\#Publication\Spacetime_Mechanics__git\docx_latex_comparison.json"
    
    print("Extracting Word document content...")
    docx_data = extract_docx_content(docx_path)
    
    print("Extracting LaTeX manuscript structure...")
    latex_data = extract_latex_structure(latex_path)
    
    print("Comparing documents...")
    comparison = compare_documents(docx_data, latex_data)
    
    # Save results
    output = {
        'docx_data': docx_data,
        'latex_data': latex_data,
        'comparison': comparison,
    }
    
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(output, f, indent=2, ensure_ascii=False)
    
    print(f"\n=== EXTRACTION COMPLETE ===")
    print(f"\nWord Document (v5.1) Statistics:")
    print(f"  - Total words: {docx_data['metadata']['total_words']:,}")
    print(f"  - Total paragraphs: {docx_data['metadata']['total_paragraphs']:,}")
    print(f"  - Sections: {len(docx_data['sections'])}")
    print(f"  - Tables: {len(docx_data['tables'])}")
    
    print(f"\nLaTeX Manuscript Statistics:")
    print(f"  - Total words: {latex_data['metadata']['total_words']:,}")
    print(f"  - Total lines: {latex_data['metadata']['total_lines']:,}")
    print(f"  - Sections: {len(latex_data['sections'])}")
    print(f"  - Unique citations: {len(set(latex_data['citations']))}")
    
    print(f"\nComparison:")
    print(f"  - New sections in v5.1: {len(comparison['new_sections'])}")
    print(f"  - Removed sections: {len(comparison['removed_sections'])}")
    
    print(f"\nOutput saved to: {output_path}")
    
    # Print section details
    print("\n=== SECTION STRUCTURE (v5.1) ===")
    for i, section in enumerate(docx_data['sections'], 1):
        print(f"{i}. [{section['level']}] {section['title']}")
        if 'subsections' in section:
            for sub in section['subsections']:
                print(f"   [{sub['level']}] {sub['title']}")

if __name__ == '__main__':
    main()
